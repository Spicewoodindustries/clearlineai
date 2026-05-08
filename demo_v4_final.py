import os
import json
import base64
import re
import io
import requests
from datetime import datetime
from dotenv import load_dotenv
import anthropic
import pypdf
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from msal import ConfidentialClientApplication

load_dotenv()

CLIENT_ID     = os.getenv("AZURE_CLIENT_ID")
TENANT_ID     = os.getenv("AZURE_TENANT_ID")
CLIENT_SECRET = os.getenv("AZURE_CLIENT_SECRET")
ANTHROPIC_KEY = os.getenv("ANTHROPIC_API_KEY")
YOUR_EMAIL    = os.getenv("BROKER_EMAIL")

GRAPH_BASE = "https://graph.microsoft.com/v1.0"

MARKETS = [
    { "name": "Intact Insurance",    "email": "commercial.quotes@intact.ca",      "lines": ["CGL", "Umbrella"] },
    { "name": "Aviva Canada",        "email": "newbusiness@aviva.ca",              "lines": ["CGL", "Property"] },
    { "name": "Wawanesa Mutual",     "email": "commercial@wawanesa.com",           "lines": ["CGL"] },
    { "name": "Travelers Canada",    "email": "submissions@travelers.ca",          "lines": ["CGL", "Umbrella"] },
    { "name": "Northbridge",         "email": "underwriting@nbins.com",            "lines": ["CGL"] },
    { "name": "FM Global",           "email": "canada.submissions@fmglobal.com",   "lines": ["Property"] },
    { "name": "Zurich Canada",       "email": "property.ca@zurich.com",            "lines": ["Property"] },
    { "name": "RSA Canada",          "email": "commercial@rsagroup.ca",            "lines": ["Property", "CGL"] },
    { "name": "Lloyd's of London",   "email": "canada.excess@lloyds.com",          "lines": ["Umbrella"] },
    { "name": "Everest Re Canada",   "email": "casualty@everestre.ca",             "lines": ["Umbrella"] },
]


# ── STEP 1: AUTHENTICATE ───────────────────────────────────────────────────────

def get_access_token():
    print("→ Authenticating with Microsoft...")
    app = ConfidentialClientApplication(
        client_id=CLIENT_ID,
        client_credential=CLIENT_SECRET,
        authority=f"https://login.microsoftonline.com/{TENANT_ID}"
    )
    result = app.acquire_token_for_client(
        scopes=["https://graph.microsoft.com/.default"]
    )
    if "access_token" not in result:
        raise Exception(f"Auth failed: {result.get('error_description', result)}")
    print("  ✓ Authenticated")
    return result["access_token"]


# ── STEP 2: READ EMAIL + ATTACHMENTS ──────────────────────────────────────────

def get_latest_email(token):
    print("→ Fetching most recent email from inbox...")
    headers = {"Authorization": f"Bearer {token}"}
    url = f"{GRAPH_BASE}/users/{YOUR_EMAIL}/messages"
    params = {
        "$orderby": "receivedDateTime desc",
        "$top": 1,
        "$select": "id,subject,body,from,receivedDateTime,hasAttachments"
    }
    resp = requests.get(url, headers=headers, params=params)
    resp.raise_for_status()
    messages = resp.json().get("value", [])
    if not messages:
        raise Exception("No emails found in inbox.")
    email = messages[0]
    body_html  = email["body"]["content"]
    body_clean = body_html.replace("<br>", "\n").replace("<p>", "\n").replace("</p>", "")
    body_text  = re.sub(r'<[^>]+>', '', body_clean).strip()
    print(f"  ✓ Found: '{email['subject']}' from {email['from']['emailAddress']['address']}")
    print(f"  ✓ Has attachments: {email.get('hasAttachments', False)}")
    return {
        "id":              email["id"],
        "subject":         email["subject"],
        "body":            body_text,
        "has_attachments": email.get("hasAttachments", False)
    }


def get_attachments(token, message_id):
    print("→ Downloading attachments...")
    headers = {"Authorization": f"Bearer {token}"}
    resp = requests.get(
        f"{GRAPH_BASE}/users/{YOUR_EMAIL}/messages/{message_id}/attachments",
        headers=headers
    )
    resp.raise_for_status()
    attachments = []
    for att in resp.json().get("value", []):
        if att.get("@odata.type") != "#microsoft.graph.fileAttachment":
            continue
        name = att.get("name", "attachment")
        ext  = os.path.splitext(name)[1].lower()
        if ext not in [".pdf", ".xlsx", ".xls"]:
            continue
        content = base64.b64decode(att["contentBytes"])
        attachments.append({"name": name, "extension": ext, "content": content})
        print(f"  ✓ Downloaded: {name}")
    if not attachments:
        print("  ! No PDF or Excel attachments found.")
    return attachments


# ── STEP 3: READ DOCUMENTS ────────────────────────────────────────────────────

def read_pdf(content_bytes, filename):
    reader = pypdf.PdfReader(io.BytesIO(content_bytes))
    parts  = []
    fields = reader.get_fields()
    if fields:
        parts.append(f"[FORM FIELDS — {filename}]")
        for name, obj in fields.items():
            value = obj.get("/V", "")
            if value and str(value).strip():
                parts.append(f"{name.strip()}: {str(value).strip()}")
    parts.append(f"[PAGE TEXT — {filename}]")
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def read_excel(content_bytes, filename):
    wb    = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
    parts = [f"[EXCEL DATA — {filename}]"]
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"\n--- Sheet: {sheet} ---")
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                parts.append(" | ".join(str(c) if c is not None else "" for c in row))
    return "\n".join(parts)


def read_all_documents(email_body, attachments):
    parts = []
    if email_body.strip():
        parts.append("[EMAIL BODY]\n" + email_body)
    for att in attachments:
        try:
            if att["extension"] == ".pdf":
                text = read_pdf(att["content"], att["name"])
                parts.append(text)
                print(f"  ✓ Read PDF: {att['name']}")
            elif att["extension"] in [".xlsx", ".xls"]:
                text = read_excel(att["content"], att["name"])
                parts.append(text)
                print(f"  ✓ Read Excel: {att['name']}")
        except Exception as e:
            print(f"  ! Could not read {att['name']}: {e}")
    return "\n\n" + ("=" * 60) + "\n\n".join(parts)


# ── STEP 4: EXTRACT WITH CLAUDE ───────────────────────────────────────────────

def extract_submission_data(combined_text):
    print("→ Extracting submission data with Claude...")
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
    prompt = f"""You are an expert commercial insurance submission analyst.
You have been given text extracted from an insurance submission — this may include
a broker email, filled PDF application forms, loss run spreadsheets, and property schedules.

Extract all relevant underwriting information and return a single JSON object.
If a field is not found, use "Not provided". Return ONLY valid JSON, no other text.

{{
  "insured_name": "",
  "insured_address": "",
  "year_established": "",
  "annual_revenue": "",
  "num_employees": "",
  "operations_description": "",
  "coverages_requested": [{{"line": "", "limits": ""}}],
  "prior_losses": [{{"year": "", "description": "", "amount": ""}}],
  "expiry_date": "",
  "incumbent_carrier": "",
  "reason_for_marketing": "",
  "submitting_broker": "",
  "property_locations": [{{"address": "", "building_value": "", "contents_value": "", "construction": "", "year_built": ""}}],
  "special_notes": "",
  "missing_information": []
}}

For missing_information, list any fields an underwriter would expect that are blank or incomplete.

Documents:
{combined_text}"""

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = message.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    data = json.loads(raw.strip())
    print(f"  ✓ Extracted data for: {data.get('insured_name', 'Unknown')}")
    return data


# ── STEP 5: MARKET SELECTION ──────────────────────────────────────────────────

def select_markets():
    print("\n" + "─" * 60)
    print("  SELECT MARKETS TO APPROACH")
    print("─" * 60)
    for i, m in enumerate(MARKETS):
        lines = ", ".join(m["lines"])
        print(f"  [{i+1:2d}]  {m['name']:<25}  {lines}")
    print("─" * 60)
    print("  Enter market numbers separated by commas (e.g. 1,3,5)")
    print("  Or type 'all' to select all markets")
    print("─" * 60)
    while True:
        raw = input("  Your selection: ").strip()
        if not raw:
            print("  Please enter at least one number.")
            continue
        if raw.lower() == "all":
            return MARKETS[:]
        try:
            indices = [int(x.strip()) - 1 for x in raw.split(",")]
            if any(i < 0 or i >= len(MARKETS) for i in indices):
                print(f"  Please enter numbers between 1 and {len(MARKETS)}.")
                continue
            selected = [MARKETS[i] for i in indices]
            print(f"\n  ✓ Selected {len(selected)} market(s):")
            for m in selected:
                print(f"    · {m['name']}")
            print()
            return selected
        except ValueError:
            print("  Invalid input — use numbers separated by commas.")


# ── STEP 6: BUILD WORD DOC ────────────────────────────────────────────────────

def build_submission_doc(data, selected_markets):
    print("→ Building submission document...")
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    TEAL  = RGBColor(0, 197, 161)
    NAVY  = RGBColor(13, 31, 60)
    SLATE = RGBColor(74, 96, 128)
    MUTED = RGBColor(138, 160, 184)
    RED   = RGBColor(192, 0, 0)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def add_teal_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '00C5A1')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL
        run.font.name = "Calibri"
        add_teal_border(p)

    def field(label, value):
        if not value or str(value).strip() in ["", "Not provided"]:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY
        r1.font.name = "Calibri"
        r2 = p.add_run(str(value))
        r2.font.size = Pt(10)
        r2.font.color.rgb = SLATE
        r2.font.name = "Calibri"

    def body(text, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = color if color else SLATE
        r.font.name = "Calibri"

    # Title block
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("CLEARLINE AI")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = TEAL; r.font.name = "Calibri"

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("COMMERCIAL INSURANCE SUBMISSION")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY; r.font.name = "Calibri"

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    r.font.size = Pt(10); r.font.color.rgb = MUTED; r.font.name = "Calibri"
    doc.add_paragraph()

    section_heading("1. Insured Information")
    field("Insured Name",     data.get("insured_name"))
    field("Address",          data.get("insured_address"))
    field("Year Established", data.get("year_established"))
    field("Annual Revenue",   data.get("annual_revenue"))
    field("Employees",        data.get("num_employees"))
    doc.add_paragraph()

    section_heading("2. Operations")
    body(data.get("operations_description", "Not provided"))
    doc.add_paragraph()

    section_heading("3. Coverages Requested")
    for cov in data.get("coverages_requested", []):
        if cov.get("line"):
            body(f"• {cov.get('line', '')}: {cov.get('limits', '')}")
    doc.add_paragraph()

    locations = data.get("property_locations", [])
    if locations and any(loc.get("address") for loc in locations):
        section_heading("4. Property Locations")
        for i, loc in enumerate(locations, 1):
            if not loc.get("address"):
                continue
            p = doc.add_paragraph()
            r = p.add_run(f"Location {i}")
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY; r.font.name = "Calibri"
            field("  Address",        loc.get("address"))
            field("  Building Value", loc.get("building_value"))
            field("  Contents Value", loc.get("contents_value"))
            field("  Construction",   loc.get("construction"))
            field("  Year Built",     loc.get("year_built"))
        doc.add_paragraph()

    section_heading("5. Loss History (5 Years)")
    losses = data.get("prior_losses", [])
    if losses and losses[0].get("year"):
        for loss in losses:
            body(f"• {loss.get('year', '')} — {loss.get('description', '')} [{loss.get('amount', '')}]")
    else:
        body("No prior losses reported.")
    doc.add_paragraph()

    section_heading("6. Market Information")
    field("Expiry Date",          data.get("expiry_date"))
    field("Incumbent Carrier",    data.get("incumbent_carrier"))
    field("Reason for Marketing", data.get("reason_for_marketing"))
    field("Submitting Broker",    data.get("submitting_broker"))
    field("Markets Approached",   ", ".join(m["name"] for m in selected_markets))
    doc.add_paragraph()

    notes = data.get("special_notes", "")
    if notes and notes != "Not provided":
        section_heading("7. Special Notes")
        body(notes)
        doc.add_paragraph()

    missing = data.get("missing_information", [])
    if missing:
        section_heading("⚠  Missing Information — Follow Up Required")
        for item in missing:
            body(f"• {item}", color=RED)
        doc.add_paragraph()

    f_para = doc.add_paragraph()
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f_para.add_run("Generated by Clearline AI  |  clearlineai.ca  |  Confidential")
    r.font.size = Pt(8); r.font.color.rgb = MUTED; r.font.name = "Calibri"

    insured_slug = data.get("insured_name", "Unknown").replace(" ", "_").replace("/", "-")
    filename = f"Submission_{insured_slug}_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    print(f"  ✓ Saved: {filename}")
    return filename


# ── STEP 7: CREATE OUTLOOK DRAFT ──────────────────────────────────────────────

def create_outlook_draft(token, data, doc_filename, selected_markets):
    print("→ Creating draft email in Outlook...")
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

    insured   = data.get("insured_name", "Unknown Insured")
    expiry    = data.get("expiry_date", "TBD")
    revenue   = data.get("annual_revenue", "See attached")
    ops       = data.get("operations_description", "See attached.")
    coverages = data.get("coverages_requested", [])
    cov_lines = "\n".join([f"  • {c['line']}: {c['limits']}" for c in coverages if c.get("line")])

    subject = f"SUBMISSION – {insured} – Expiry {expiry}"
    body_text = f"""Dear Underwriter,

Please find attached a commercial insurance submission for your review and quotation.

INSURED:           {insured}
EXPIRY DATE:       {expiry}
ANNUAL REVENUE:    {revenue}

OPERATIONS SUMMARY:
{ops}

COVERAGES REQUESTED:
{cov_lines}

A complete submission package is attached including full risk details, loss history, and underwriting notes.

Please confirm receipt and advise your earliest indication.

Best regards,
[Your Name]
[Your Brokerage]
clearlineai.ca
"""

    with open(doc_filename, "rb") as f:
        doc_b64 = base64.b64encode(f.read()).decode("utf-8")

    to_recipients = [
        {"emailAddress": {"address": m["email"], "name": m["name"]}}
        for m in selected_markets
    ]

    draft_payload = {
        "subject": subject,
        "body": {"contentType": "Text", "content": body_text},
        "toRecipients": to_recipients,
        "attachments": [{
            "@odata.type": "#microsoft.graph.fileAttachment",
            "name": doc_filename,
            "contentBytes": doc_b64,
            "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }]
    }

    url = f"{GRAPH_BASE}/users/{YOUR_EMAIL}/messages"
    resp = requests.post(url, headers=headers, json=draft_payload)
    resp.raise_for_status()

    print(f"  ✓ Draft created in Outlook Drafts")
    print(f"  → Subject: {subject}")
    print(f"  → To: {', '.join(m['name'] for m in selected_markets)}")
    print(f"  → Attachment: {doc_filename}")


# ── MAIN ──────────────────────────────────────────────────────────────────────

def main():
    print("\n" + "=" * 60)
    print("  CLEARLINE AI — SUBMISSION AUTOMATOR")
    print("=" * 60 + "\n")

    start = datetime.now()

    token = get_access_token()
    email = get_latest_email(token)

    attachments = []
    if email["has_attachments"]:
        attachments = get_attachments(token, email["id"])

    if not attachments:
        print("\n  ! No attachments found — processing email body text only.")
        print("  Tip: Forward a client submission with PDF/Excel files attached.\n")

    print("→ Reading documents...")
    combined_text = read_all_documents(email["body"], attachments)

    data    = extract_submission_data(combined_text)
    markets = select_markets()
    docfile = build_submission_doc(data, markets)
    create_outlook_draft(token, data, docfile, markets)

    elapsed = (datetime.now() - start).seconds
    print(f"\n{'=' * 60}")
    print(f"  ✓ COMPLETE in {elapsed} seconds")
    print(f"  Document : {docfile}")
    print(f"  Draft email is in your Outlook Drafts — ready to send.")
    print(f"{'=' * 60}\n")


if __name__ == "__main__":
    main()
